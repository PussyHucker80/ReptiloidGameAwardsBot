# -*- coding: utf-8 -*-
"""
Game Awards Telegram Bot
Автор: (пример)
Требует: python-telegram-bot==20.x, pandas, python-docx, openpyxl, reportlab
Запуск:
    set TELEGRAM_TOKEN=8413467526:AAFukjD4IkPniFbFBRiW5mCip_gpeLIoZNk     (Windows cmd)
    # или в PowerShell:
    $env:TELEGRAM_TOKEN="8413467526:AAFukjD4IkPniFbFBRiW5mCip_gpeLIoZNk"
    python game_awards_bot.py
"""

import os
import sqlite3
import json
import logging
from datetime import datetime
from typing import Dict, List, Tuple

from telegram import (
    Update,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    Poll,
)
from telegram.ext import (
    ApplicationBuilder,
    ContextTypes,
    CommandHandler,
    CallbackQueryHandler,
    MessageHandler,
    filters,
)

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# ------- Настройки -------
TOKEN = os.getenv("TELEGRAM_TOKEN")  # берём из переменных окружения
DB_PATH = "game_awards.db"
ADMIN_USER_IDS = []  # сюда можно записать Telegram user_id админов (опционально). Если пустой - команду может выполнять любой, кто является creator/админ чата (проверяется динамически)
MAX_POLL_OPTIONS = 10  # Telegram лимит опций в одном poll
# -------------------------

# Логирование
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO
)
logger = logging.getLogger(__name__)


# ---------- База данных (very small ORM over sqlite3) ----------
def init_db():
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS categories(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT UNIQUE,
        created_by INTEGER,
        created_by_name TEXT,
        created_at TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS games(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        category_id INTEGER,
        suggested_by INTEGER,
        suggested_by_name TEXT,
        suggested_at TEXT,
        UNIQUE(title, category_id)
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS polls(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        telegram_poll_id TEXT,
        category_id INTEGER,
        options_json TEXT,
        active INTEGER DEFAULT 1,
        created_at TEXT
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS votes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        telegram_poll_id TEXT,
        telegram_message_id INTEGER,
        user_id INTEGER,
        username TEXT,
        game_id INTEGER,
        option_index INTEGER,
        voted_at TEXT
    )
    """)
    con.commit()
    con.close()


def db_execute(query: str, params=(), fetch=False, many=False):
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    if many:
        cur.executemany(query, params)
        res = None
    else:
        cur.execute(query, params)
        res = cur.fetchall() if fetch else None
    con.commit()
    con.close()
    return res


# ---------- Утилиты работы с БД ----------
def add_category(title: str, user_id: int, user_name: str):
    now = datetime.utcnow().isoformat()
    try:
        db_execute(
            "INSERT INTO categories(title, created_by, created_by_name, created_at) VALUES (?, ?, ?, ?)",
            (title, user_id, user_name, now),
        )
        return True
    except sqlite3.IntegrityError:
        return False


def list_categories() -> List[Tuple]:
    res = db_execute("SELECT id, title FROM categories ORDER BY id", fetch=True)
    return res or []


def add_game(title: str, category_id: int, user_id: int, user_name: str):
    now = datetime.utcnow().isoformat()
    try:
        db_execute(
            "INSERT INTO games(title, category_id, suggested_by, suggested_by_name, suggested_at) VALUES (?, ?, ?, ?, ?)",
            (title, category_id, user_id, user_name, now),
        )
        return True
    except sqlite3.IntegrityError:
        return False


def list_games_for_category(category_id: int) -> List[Tuple]:
    res = db_execute(
        "SELECT id, title, suggested_by_name FROM games WHERE category_id = ? ORDER BY id",
        (category_id,),
        fetch=True,
    )
    return res or []


def store_poll(telegram_poll_id: str, category_id: int, options_map: Dict[int, int]):
    # options_map: option_index -> game_id
    now = datetime.utcnow().isoformat()
    db_execute(
        "INSERT INTO polls(telegram_poll_id, category_id, options_json, created_at) VALUES (?, ?, ?, ?)",
        (telegram_poll_id, category_id, json.dumps(options_map), now),
    )


def mark_poll_closed(telegram_poll_id: str):
    db_execute("UPDATE polls SET active = 0 WHERE telegram_poll_id = ?", (telegram_poll_id,))


def get_poll_by_tg_id(telegram_poll_id: str):
    res = db_execute("SELECT id, options_json, category_id, active FROM polls WHERE telegram_poll_id = ?", (telegram_poll_id,), fetch=True)
    return res[0] if res else None


def record_vote(telegram_poll_id: str, telegram_message_id: int, user_id: int, username: str, game_id: int, option_index: int):
    now = datetime.utcnow().isoformat()
    db_execute(
        "INSERT INTO votes(telegram_poll_id, telegram_message_id, user_id, username, game_id, option_index, voted_at) VALUES (?, ?, ?, ?, ?, ?, ?)",
        (telegram_poll_id, telegram_message_id, user_id, username, game_id, option_index, now),
    )


def delete_votes_for_poll(telegram_poll_id: str):
    db_execute("DELETE FROM votes WHERE telegram_poll_id = ?", (telegram_poll_id,))


def list_votes_for_category(category_id: int):
    res = db_execute("""
    SELECT v.user_id, v.username, g.title, c.title, v.voted_at
    FROM votes v
    LEFT JOIN games g ON v.game_id = g.id
    LEFT JOIN categories c ON g.category_id = c.id
    WHERE g.category_id = ?
    ORDER BY v.voted_at
    """, (category_id,), fetch=True)
    return res or []


def list_all_votes():
    res = db_execute("""
    SELECT c.title AS category, g.title AS game, v.username, v.user_id, v.voted_at
    FROM votes v
    LEFT JOIN games g ON v.game_id = g.id
    LEFT JOIN categories c ON g.category_id = c.id
    ORDER BY c.id, g.id
    """, fetch=True)
    return res or []

# ---------- Telegram handlers & logic ----------

# Стейты на время вводов (in-memory, per chat)
# Для простоты: словарь chat_id -> state dict
chat_states = {}


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Команда /start — приватно или в чате"""
    text = (
        "Привет! Я бот для голосования — помогу собрать предложения игр и провести НЕ-анонимное голосование.\n\n"
        "Кнопки:\n"
        "- Предложить игру\n"
        "- Предложить категорию\n"
        "- Создать голосование (только для админа/создателя чата)\n"
        "- Закрыть голосование (только для админа)\n"
        "- Экспорт результатов (только для админа)"
    )
    keyboard = [
        [InlineKeyboardButton("📥 Предложить игру", callback_data="suggest_game")],
        [InlineKeyboardButton("➕ Предложить категорию", callback_data="suggest_category")],
        [InlineKeyboardButton("📊 Создать голосование", callback_data="create_poll")],
        [InlineKeyboardButton("🔒 Закрыть голосование", callback_data="close_poll")],
        [InlineKeyboardButton("📤 Экспорт данных", callback_data="export_data")],
    ]
    await update.effective_chat.send_message(text, reply_markup=InlineKeyboardMarkup(keyboard))


async def button_router(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Маршрутизация нажатий кнопок"""
    query = update.callback_query
    await query.answer()
    data = query.data
    chat_id = query.message.chat_id
    user = query.from_user

    # Инициализируем стейт для чата
    chat_states.setdefault(chat_id, {})

    if data == "suggest_game":
        # Выбираем категорию
        cats = list_categories()
        if not cats:
            await query.message.reply_text("Пока нет категорий. Попросите добавить категорию (➕ Предложить категорию).")
            return
        kb = []
        for cid, title in cats:
            kb.append([InlineKeyboardButton(title, callback_data=f"suggest_game_cat:{cid}")])
        await query.message.reply_text("Выберите категорию для предложения игры:", reply_markup=InlineKeyboardMarkup(kb))

    elif data.startswith("suggest_game_cat:"):
        cat_id = int(data.split(":", 1)[1])
        chat_states[chat_id]['awaiting_game_for_cat'] = cat_id
        await query.message.reply_text("Введите название игры (на английском желателен):")

    elif data == "suggest_category":
        chat_states[chat_id]['awaiting_new_category'] = True
        await query.message.reply_text("Введите название новой категории (например: 'Лучшая RPG 2025 года'):")

    elif data == "create_poll":
        # Только админы чата (или если ADMIN_USER_IDS содержит id)
        if not await user_is_admin_in_chat(context, chat_id, user.id):
            await query.message.reply_text("Только админ/создатель чата может создавать голосование.")
            return
        cats = list_categories()
        if not cats:
            await query.message.reply_text("Нет категорий для голосования. Добавьте хотя бы одну категорию.")
            return
        kb = []
        for cid, title in cats:
            kb.append([InlineKeyboardButton(title, callback_data=f"create_poll_cat:{cid}")])
        await query.message.reply_text("Выберите категорию для создания голосования:", reply_markup=InlineKeyboardMarkup(kb))

    elif data.startswith("create_poll_cat:"):
        cat_id = int(data.split(":",1)[1])
        # Получаем кандидатов
        games = list_games_for_category(cat_id)
        if not games:
            await query.message.reply_text("В этой категории нет предложенных игр. Попросите участников предложить игры.")
            return
        # Формируем опции (учитывая лимит)
        options = [g[1] for g in games]  # title
        # если больше MAX_POLL_OPTIONS - делим на несколько опросов (пагинация)
        chunks = [options[i:i+MAX_POLL_OPTIONS] for i in range(0, len(options), MAX_POLL_OPTIONS)]
        # Мы отправляем серию опросов и запоминаем mapping опция->game_id для каждого poll
        sent_info = []
        for idx, chunk in enumerate(chunks, start=1):
            # find corresponding game ids for chunk
            option_to_gameid = {}
            for opt_index, opt_text in enumerate(chunk):
                # find game_id
                for g in games:
                    if g[1] == opt_text:
                        option_to_gameid[opt_index] = g[0]
                        break
            # текст заголовка
            title_text = f"Голосование — {get_category_title(cat_id)}"
            if len(chunks) > 1:
                title_text += f" (часть {idx}/{len(chunks)})"
            # отправляем опрос (не анонимный)
            message = await context.bot.send_poll(
                chat_id=chat_id,
                question=title_text,
                options=chunk,
                is_anonymous=False,
                allows_multiple_answers=False,
            )
            # Сохраняем mapping между telegram_poll.id и game ids
            tg_poll_id = message.poll.id  # уникальный идентификатор опроса
            # map option_index -> game_id
            store_poll(tg_poll_id, cat_id, option_to_gameid)
            sent_info.append((tg_poll_id, message.message_id))
        await query.message.reply_text(f"Отправлено {len(sent_info)} опрос(ов) для категории '{get_category_title(cat_id)}'. Голосование активно.")

    elif data == "close_poll":
        if not await user_is_admin_in_chat(context, chat_id, user.id):
            await query.message.reply_text("Только админ/создатель чата может закрыть голосование.")
            return
        # Предложим список активных опросов из polls
        res = db_execute("SELECT id, telegram_poll_id FROM polls WHERE active=1", fetch=True)
        if not res:
            await query.message.reply_text("Нет активных опросов для закрытия.")
            return
        kb = [[InlineKeyboardButton(f"Закрыть опрос #{row[0]}", callback_data=f"close_poll_id:{row[1]}")] for row in res]
        await query.message.reply_text("Выберите опрос, который хотите закрыть:", reply_markup=InlineKeyboardMarkup(kb))

    elif data.startswith("close_poll_id:"):
        tg_poll_id = data.split(":",1)[1]
        # закрываем: пометим active=0
        mark_poll_closed(tg_poll_id)
        await query.message.reply_text("Опрос закрыт (помечен как неактивный).")
    elif data == "export_data":
        # только админ
        if not await user_is_admin_in_chat(context, update.effective_chat.id, user.id):
            await query.message.reply_text("Только админ/создатель чата может экспортировать данные.")
            return
        # Экспорт всех голосов, игр и категорий
        await query.message.reply_text("Генерирую файлы экспорта... Подождите.")
        # Генерируем файлы
        export_folder = generate_exports()
        await query.message.reply_text(f"Файлы экспортированы в папку: {export_folder}\nФайлы: votes.xlsx, votes.docx, votes.pdf")
    else:
        await query.message.reply_text("Неопознанная команда кнопки.")


async def user_is_admin_in_chat(context: ContextTypes.DEFAULT_TYPE, chat_id: int, user_id: int) -> bool:
    """Попытка проверить является ли пользователь админом/creator в чате.
       Если ADMIN_USER_IDS задан, проверяем там (в приоритете)."""
    if ADMIN_USER_IDS:
        return user_id in ADMIN_USER_IDS
    try:
        member = await context.bot.get_chat_member(chat_id, user_id)
        status = member.status  # 'creator', 'administrator', 'member', ...
        return status in ("creator", "administrator")
    except Exception as e:
        logger.warning("Не удалось получить информацию о правах: %s", e)
        return False


def get_category_title(cat_id: int) -> str:
    res = db_execute("SELECT title FROM categories WHERE id = ?", (cat_id,), fetch=True)
    return res[0][0] if res else "Unknown"


async def text_message_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработка вводимых пользователем названий игр и категорий (простая state-machine по chat_states)"""
    msg = update.message
    chat_id = update.effective_chat.id
    user = update.effective_user
    text = msg.text.strip()
    chat_states.setdefault(chat_id, {})

    st = chat_states[chat_id]

    if st.get('awaiting_new_category'):
        added = add_category(text, user.id, user.full_name or user.username or str(user.id))
        st.pop('awaiting_new_category', None)
        if added:
            await msg.reply_text(f"Категория '{text}' добавлена.")
        else:
            await msg.reply_text(f"Категория '{text}' уже существует.")

    elif 'awaiting_game_for_cat' in st:
        cat_id = st.pop('awaiting_game_for_cat')
        ok = add_game(text, cat_id, user.id, user.full_name or user.username or str(user.id))
        if ok:
            await msg.reply_text(f"Игра '{text}' предложена в категории '{get_category_title(cat_id)}'.")
        else:
            await msg.reply_text(f"Игра '{text}' уже есть в этой категории.")
    else:
        # Если сообщение не в контексте — короткий хелп
        await msg.reply_text("Если хотите предложить игру или категорию — используйте /start и кнопки.")


async def poll_answer_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработка ответов на опросы — Telegram присылает PollAnswer, содержащий user и выбранные option_ids.
       Мы сохраняем в таблицу votes явный user->game mapping (на основе ранее сохранённого polls.options_json).
    """
    answer = update.poll_answer  # telegram.PollAnswer
    user = answer.user
    tg_poll_id = answer.poll_id
    chosen = answer.option_ids  # list of option indexes (0-based)
    # получить mapping для этого poll
    pollrow = get_poll_by_tg_id(tg_poll_id)
    if not pollrow:
        logger.info("Получен ответ на неизвестный опрос %s", tg_poll_id)
        return
    poll_db_id, options_json, category_id, active = pollrow
    options_map = json.loads(options_json)  # option_index -> game_id

    # NOTE: update.poll_answer не содержит message_id; мы не знаем message_id здесь.
    # Но мы можем записать telegram_poll_id и user -> game
    # Для simplicity: запишем все выбранные варианты (Telegram может отправлять несколько если allows_multiple_answers=True).
    # Удалим предыдущие голоса этого пользователя в этом poll (на случай изменения выбора)
    db_execute("DELETE FROM votes WHERE telegram_poll_id = ? AND user_id = ?", (tg_poll_id, user.id))

    for opt_idx in chosen:
        game_id = options_map.get(str(opt_idx)) if isinstance(options_map.keys().__iter__().__next__(), str) else options_map.get(opt_idx)
        # options_map might be stored with integer keys or string keys - normalize
        if game_id is None:
            # try string key
            game_id = options_map.get(str(opt_idx))
        if game_id is None:
            logger.warning("Не могу найти game_id для option %s in poll %s", opt_idx, tg_poll_id)
            continue
        # message_id unknown, set None (0)
        record_vote(tg_poll_id, 0, user.id, user.full_name or user.username or str(user.id), game_id, opt_idx)
    logger.info("Recorded vote(s) for user %s in poll %s", user.id, tg_poll_id)


# ---------- Export функций ----------
def generate_exports(folder="exports"):
    os.makedirs(folder, exist_ok=True)
    # Excel
    rows = list_all_votes()
    if rows:
        df = pd.DataFrame(rows, columns=["Category", "Game", "Username", "UserID", "VotedAt"])
    else:
        df = pd.DataFrame(columns=["Category", "Game", "Username", "UserID", "VotedAt"])
    xlsx_path = os.path.join(folder, "votes.xlsx")
    df.to_excel(xlsx_path, index=False)

    # Word (docx)
    doc = Document()
    doc.add_heading("Game Awards — Результаты голосования", level=1)
    if rows:
        # Группируем по категории -> game -> voters
        grouped = {}
        for cat, game, username, uid, voted_at in rows:
            grouped.setdefault(cat, {}).setdefault(game, []).append((username, voted_at))
        for cat, games in grouped.items():
            doc.add_heading(cat, level=2)
            for game, voters in games.items():
                doc.add_paragraph(f"{game} — {len(voters)} голосов")
                for username, voted_at in voters:
                    doc.add_paragraph(f" - {username} ({voted_at})", style='List Bullet')
    else:
        doc.add_paragraph("Пока нет голосов.")
    docx_path = os.path.join(folder, "votes.docx")
    doc.save(docx_path)

    # PDF (простая таблица)
    pdf_path = os.path.join(folder, "votes.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Game Awards — Результаты голосования")
    y -= 30
    c.setFont("Helvetica", 10)
    if rows:
        # печатаем шапку
        c.drawString(50, y, "Категория")
        c.drawString(220, y, "Игра")
        c.drawString(380, y, "Пользователь")
        c.drawString(520, y, "Время")
        y -= 15
        for r in rows:
            if y < 60:
                c.showPage()
                y = height - 50
            cat, game, username, uid, voted_at = r
            c.drawString(50, y, str(cat)[:30])
            c.drawString(220, y, str(game)[:30])
            c.drawString(380, y, str(username)[:20])
            c.drawString(520, y, str(voted_at)[:16])
            y -= 12
    else:
        c.drawString(50, y, "Пока нет голосов.")
    c.save()

    return os.path.abspath(folder)


# ---------- Хелп команды ----------
async def list_categories_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    cats = list_categories()
    if not cats:
        await update.message.reply_text("Категорий пока нет.")
        return
    text = "Категории:\n" + "\n".join([f"{cid}. {title}" for cid, title in cats])
    await update.message.reply_text(text)


async def list_games_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Если вызов /list_games <category_id>
    args = context.args
    if not args:
        await update.message.reply_text("Использование: /list_games <category_id>")
        return
    try:
        cat_id = int(args[0])
    except:
        await update.message.reply_text("Неверный id категории.")
        return
    games = list_games_for_category(cat_id)
    if not games:
        await update.message.reply_text("В этой категории пока нет игр.")
        return
    text = f"Игры в категории {get_category_title(cat_id)}:\n" + "\n".join([f"{gid}. {title} (предложил: {who})" for gid, title, who in games])
    await update.message.reply_text(text)


# ---------- Main ----------
def main():
    if not TOKEN:
        print("Ошибка: TELEGRAM_TOKEN не задан. Установите переменную окружения TELEGRAM_TOKEN.")
        return

    init_db()

   app = Application.builder().token(TOKEN).build()

    # Handlers
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("list_categories", list_categories_cmd))
    app.add_handler(CommandHandler("list_games", list_games_cmd))
    app.add_handler(CallbackQueryHandler(button_router))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, text_message_handler))
    app.add_handler(MessageHandler(filters.UpdateType.POLL_ANSWER, lambda u, c: None))  # placeholder
    # PollAnswer обработчик: нужно использовать специальный обработчик через .add_handler, но PTB требует PollAnswerHandler — мы используем фильтр UpdateType.POLL_ANSWER
    from telegram.ext import PollAnswerHandler
    app.add_handler(PollAnswerHandler(poll_answer_handler))

    print("Bot started...")
    app.run_polling()


if __name__ == "__main__":
    main()
